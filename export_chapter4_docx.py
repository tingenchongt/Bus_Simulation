"""
Build THESIS_Chapter4.docx from THESIS_DRAFT.md (Chapter 4 content only).
Run: python export_chapter4_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_inline_to_paragraph(p, text: str) -> None:
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            p.add_run(part)


def add_body_paragraph(doc: Document, text: str) -> None:
    if "*End of Chapter 4 draft.*" in text:
        return
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    add_inline_to_paragraph(p, text)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and "|" in s[1:]


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().split("|")[1:-1]]


def is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(re.match(r"^:?-+$", re.sub(r"-+", "-", c.replace(" ", ""))) for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if len(rows) >= 2 and is_separator_row(rows[1]):
        data_rows = [rows[0]] + rows[2:]
    else:
        data_rows = rows
    ncols = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(data_rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            p = table.rows[ri].cells[j].paragraphs[0]
            p.clear()
            add_inline_to_paragraph(p, cell_text)


def main() -> None:
    root = Path(__file__).resolve().parent
    md_path = root / "THESIS_DRAFT.md"
    out_path = root / "THESIS_Chapter4.docx"

    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    i = 0
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        if i < len(lines):
            i += 1

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if is_table_row(line):
            block: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                block.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, block)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_to_paragraph(p, stripped[2:].strip())
        elif not stripped:
            pass
        else:
            add_body_paragraph(doc, stripped)
        i += 1

    doc.core_properties.title = "Thesis Chapter 4 - UXsim Simulation"
    try:
        doc.save(out_path)
        print(f"Wrote: {out_path}")
    except PermissionError:
        alt = root / "THESIS_Chapter4_generated.docx"
        doc.save(alt)
        print(f"Could not overwrite {out_path} (close Word if open). Wrote: {alt}")


if __name__ == "__main__":
    main()
